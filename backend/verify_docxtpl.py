from docxtpl import DocxTemplate
from docx import Document
import os

# 1. Create a dummy template using python-docx
template_path = "temp_template.docx"
output_path = "temp_output.docx"

doc = Document()
doc.add_paragraph("Hello {{ name }}, welcome to {{ place }}.")
doc.save(template_path)

print(f"Created template at {template_path}")

try:
    # 2. Render using DocxTemplate
    doc_tpl = DocxTemplate(template_path)
    context = {
        "name": "Antigravity",
        "place": "Deepmind"
    }
    doc_tpl.render(context)
    doc_tpl.save(output_path)
    print(f"Rendered document saved to {output_path}")

    # 3. Verify content
    final_doc = Document(output_path)
    full_text = "\n".join([p.text for p in final_doc.paragraphs])
    print(f"Final Text: {full_text}")

    if "Hello Antigravity, welcome to Deepmind." in full_text:
        print("SUCCESS: Variables replaced correctly.")
    else:
        print("FAILURE: Text does not match expected output.")

except Exception as e:
    print(f"ERROR: {e}")

# Cleanup
if os.path.exists(template_path):
    os.remove(template_path)
if os.path.exists(output_path):
    os.remove(output_path)
