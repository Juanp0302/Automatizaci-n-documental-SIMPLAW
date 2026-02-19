import pytest
from httpx import AsyncClient
from app import crud, schemas
from sqlalchemy.orm import Session
import io

# Mock data
USER_EMAIL = "template_test@example.com"
USER_PASSWORD = "password123"

@pytest.mark.asyncio
async def test_templates_crud(client: AsyncClient, db: Session):
    # 0. Create user and login
    user_in = schemas.UserCreate(email=USER_EMAIL, password=USER_PASSWORD, full_name="Template Tester")
    user = crud.user.create(db, obj_in=user_in)
    
    login_response = await client.post(
        "/api/v1/login/access-token",
        data={"username": USER_EMAIL, "password": USER_PASSWORD},
    )
    token = login_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # 1. Create Template (Upload)
    # Create a dummy docx file in memory
    docx_content = b"PK\x03\x04\x14\x00\x00\x00\x08\x00" # Minimal zip header
    files = {"file": ("crud_template.docx", docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    
    # We need to mock the validation or ensure backend accepts this "fake" docx.
    # The backend uses python-docx.Document(file), so it might fail if content is invalid.
    # Let's generate a valid minimal one.
    from docx import Document
    doc = Document()
    doc.add_paragraph("CRUD test {{ variable }}")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    files = {"file": ("crud_template.docx", doc_io, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

    upload_response = await client.post(
        "/api/v1/templates/",
        headers=headers,
        files=files,
        data={"title": "CRUD Template", "description": "CRUD Description"}
    )
    assert upload_response.status_code == 200
    template_data = upload_response.json()
    assert template_data["title"] == "CRUD Template"
    assert template_data["description"] == "CRUD Description"
    template_id = template_data["id"]

    # 2. Get Template List
    list_response = await client.get("/api/v1/templates/", headers=headers)
    assert list_response.status_code == 200
    templates = list_response.json()
    assert len(templates) >= 1
    found = any(t["id"] == template_id for t in templates)
    assert found

    # 3. Get Single Template
    get_response = await client.get(f"/api/v1/templates/{template_id}", headers=headers)
    assert get_response.status_code == 200
    assert get_response.json()["id"] == template_id

    # 4. Delete Template
    # Verify owner can delete
    del_response = await client.delete(f"/api/v1/templates/{template_id}", headers=headers)
    assert del_response.status_code == 200
    assert del_response.json()["title"] == "CRUD Template"

    # Verify it's gone
    get_again = await client.get(f"/api/v1/templates/{template_id}", headers=headers)
    assert get_again.status_code == 404
