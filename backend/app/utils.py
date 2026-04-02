import re
import os
import json
from docx import Document
from openai import OpenAI

def get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY no configurado en las variables de entorno.")
    return OpenAI(api_key=api_key)


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        # Fallback for other potential parent types or direct body elements
        # This handles the case if parent is passed as doc.part.document
        if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
             parent_elm = parent.element.body
        else:
             raise ValueError("Could not find body for extraction")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_variables_from_docx(file_path: str) -> list[str]:
    """
    Extracts variables in the format {{variable_name}} from a .docx file.
    Returns a list of unique variable names in the order they appear.
    """
    matches_list = []
    try:
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph

        # Re-defining iter_block_items inside to avoid import issues or keeping it standalone if imports are global
        # but let's put imports inside try/except to be safe or at top level if common.
        # For simplicity in this replacement, I will assume imports are managed.
        # Actually, let's just use the logic inline or define helper properly.
        # Given the imports might not be available at module level yet.
        
        doc = Document(file_path)
        pattern = re.compile(r"\{\{([^}]+)\}\}")

        def search_in_text(text):
            found = pattern.findall(text)
            for match in found:
                matches_list.append(match.strip())

        def iter_block_items_local(parent):
            """
            Yield each paragraph and table child within *parent*, in document order.
            """
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                return # Should not happen with standard python-docx objects we use

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        def process_container(container):
            for block in iter_block_items_local(container):
                if isinstance(block, Paragraph):
                    search_in_text(block.text)
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            process_container(cell)

        # 1. Search in headers
        for section in doc.sections:
             for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    # Headers only contain paragraphs/tables usually, but python-docx header.paragraphs/tables are separate
                    # We can try to use iteration if getting element is possible, else fall back to separate
                    # Headers in python-docx don't expose 'iter_block_items' easily because they are not Document objects
                    # but they have a .part and .element.
                    # For safety/simplicity, if order CRITICALLY matters between a header paragraph and a header table, we need it.
                    # Usually headers are small. Let's stick to standard separate implementation for headers/footers 
                    # UNLESS user complains. The main issue was body text/tables mixing.
                    # Wait, the user said "variables in the configuration menu... do not appear in the same order".
                    # This implies body content.
                    for paragraph in header.paragraphs:
                        search_in_text(paragraph.text)
                    for table in header.tables:
                        for row in table.rows:
                             for cell in row.cells:
                                 for paragraph in cell.paragraphs:
                                     search_in_text(paragraph.text)

        # 2. Search in main body (Document)
        process_container(doc)

        # 3. Search in footers
        for section in doc.sections:
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        search_in_text(paragraph.text)
                    for table in footer.tables:
                        for row in table.rows:
                             for cell in row.cells:
                                 for paragraph in cell.paragraphs:
                                     search_in_text(paragraph.text)

    except Exception as e:
        print(f"Error reading docx for variables: {e}")
        import traceback
        traceback.print_exc()
        return []

    # Remove duplicates while preserving order
    return list(dict.fromkeys(matches_list))

def extract_text_from_pdf(file_path: str) -> str:
    """Extrae todo el texto de un archivo .pdf para análisis de IA."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- PÁGINA {page_num} ---")
                text_parts.append(page_text)
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        print("PyMuPDF no instalado. Instala con: pip install pymupdf")
        return ""
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """Extrae texto de un .docx o .pdf según la extensión del archivo."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        return extract_text_from_docx(file_path)

def extract_text_from_docx(file_path: str) -> str:
    """Extrae todo el texto y tablas de un archivo .docx para análisis de IA."""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
                
        for table in doc.tables:
            text.append("--- INICIO TABLA ---")
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                text.append(" | ".join(row_data))
            text.append("--- FIN TABLA ---")
            
        return "\n".join(text)
    except Exception as e:
        print(f"Error extracting text for AI analysis: {e}")
        return ""

def generate_template_proposal(doc_text: str, user_prompt: str = None) -> dict:
    """Genera una propuesta de variables usando OpenAI."""
    extra_instructions = ""
    if user_prompt and user_prompt.strip():
        extra_instructions = f"\n\nINSTRUCCIONES DEL USUARIO (MÁXIMA PRIORIDAD - SIGUE ESTAS INSTRUCCIONES ANTE TODO):\n{user_prompt}\n"

    # Si el documento está vacío pero hay prompt del usuario, intentar de todos modos
    effective_text = doc_text.strip() if doc_text else ""
    if not effective_text and not (user_prompt and user_prompt.strip()):
        return {"simple": [], "groups": []}
    if not effective_text:
        effective_text = "[Documento sin texto extraíble - usar instrucciones del usuario para generar variables]"

    prompt = f"""Eres un experto en automatización de documentos legales y comerciales en español.
Tu tarea es analizar un documento y definir su ESTRUCTURA DE VARIABLES (el formulario de preguntas que otra persona llenará después para generar el documento final).

SUPERPROMPT - REGLAS FUNDAMENTALES (NO NEGOCIABLES):
1. **OBJETIVO DE LA HERRAMIENTA**: Hacer automatizaciones a partir de plantillas. Un usuario genera una plantilla base, y luego otro usuario llena "espacios" respondiendo preguntas para generar el documento con información nueva.
2. **ESQUELETO VS DATOS**: Identifica estrictamente qué hace parte del "esqueleto" (texto estático que siempre va a estar) y qué información debe cambiar en cada documento nuevo a partir de tu conocimiento del tipo de documento.
3. **PREGUNTAS NECESARIAS**: Para cada campo dinámico, debes formular la pregunta que le harías al usuario final en el campo "label". Sé descriptivo.
4. **LISTAS Y GRUPOS DINÁMICOS (CRÍTICO)**: Si el documento tiene una lista con elementos repetidos (ej. pago_1, pago_2 o servicio_1, servicio_2) o una tabla con múltiples filas de la misma estructura, ¡NO DEBES crear variables individuales en "simple" (ni servicio_1, ni servicio_2)! DEBES crear un ÚNICO grupo en la lista "groups". Por ejemplo, crea el grupo "servicios" y adentro pon sus campos estandarizados (descripción, valor, etc.). La herramienta se encargará de preguntar cuántos elementos quiere el usuario.
5. **INFERENCIA CONTEXTUAL**: En respuestas a PQR cambia identificación, objeto, si es afirmativo o no (puedes sugerir options tipo select), razones, etc. Parametriza TODO lo que típicamente cambiaría.
6. SIEMPRE propones variables. Si el texto tiene datos reales parametrizalos. Si está casi vacío, créalas de cero basándote en la solicitud.{extra_instructions}

Ejemplo de respuesta válida estricta en JSON puro sin caracteres extra:
{{"simple": [{{"name": "nombre_cliente", "label": "¿Cuál es el nombre completo del cliente/usuario?", "type": "text"}}, {{"name": "decide_aprobar", "label": "¿La respuesta a la solicitud es afirmativa?", "type": "select", "options": ["Sí::si", "No::no"]}}, {{"name": "razones_decision", "label": "¿Cuáles son las razones de la decisión?", "type": "textarea"}}], "groups": [{{"name": "servicios", "label": "Servicios a prestar", "fields": [{{"name": "nombre_servicio", "label": "¿Nombre del servicio a prestar?", "type": "text"}}, {{"name": "actividades", "label": "¿Cuáles son las actividades que componen este servicio?", "type": "textarea"}}, {{"name": "valor_servicio", "label": "¿Valor de este servicio?", "type": "text"}}]}}]}}

Texto del documento base:
{effective_text[:8000]}
"""

    try:
        client = get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Eres un asistente de automatización JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        content = response.choices[0].message.content.strip()
        # Clean markdown codeblocks if AI happens to return them despite instructions
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
            
        return json.loads(content.strip())
    except Exception as e:
        print(f"Error en OpenAI API o parseando JSON: {e}")
        return {"simple": [], "groups": []}


def review_document_with_ai(docx_path: str) -> dict:
    """
    Revisa un documento .docx generado con IA para corregir errores de redacción
    y asegurar consistencia de datos (montos, nombres, fechas).

    Estrategia:
    1. Extrae el texto completo del .docx.
    2. Envía a GPT-4o-mini pidiendo correcciones como pares JSON: [{original, corrected}].
    3. Aplica cada par como find/replace en párrafos y celdas de tabla (conserva formato).
    4. Retorna {"corrections": N, "details": [...]} con el número de correcciones aplicadas.
    """
    try:
        # --- 1. Extraer texto completo para el análisis ---
        doc_text = extract_text_from_docx(docx_path)
        if not doc_text.strip():
            return {"corrections": 0, "details": [], "summary": "Documento vacío, sin revisión."}

        # --- 2. Obtener correcciones de la IA ---
        prompt = f"""Eres un asistente especializado en revisión de documentos legales y contractuales en español.
Se te proporciona el texto completo de un documento ya diligenciado con datos reales de un cliente.

Tu tarea es:
1. Detectar errores de redacción: concordancia de género y número, tildes, puntuación, imprecisión de lenguaje formal.
2. Detectar inconsistencias de datos: si un mismo concepto (ej. monto de honorarios, nombre de parte, fecha) aparece con valores diferentes en distintas partes del documento, señala la inconsistencia y propón el valor más apropiado (generalmente el primero que aparece o el más referenciado).
3. NO cambiar valores de variables que el usuario ingresó intencionalmente como distintos. Solo corregir cuando la inconsistencia parece un error.
4. NO alterar cláusulas legales sustantivas ni el significado de los acuerdos.
5. NO recalcular ni alterar sumas matemáticas o totales (ej. precios totales, sumatorias). Los documentos pueden listar un total correcto sin mostrar el desglose completo en el texto.
6. Si el documento está correcto, devuelve un array vacío [].

Responde ÚNICAMENTE con un array JSON válido de objetos con esta estructura EXACTA:
[
  {{"original": "texto exacto a reemplazar (como aparece en el documento)", "corrected": "texto corregido"}},
  ...
]

IMPORTANTE:
- "original" debe ser una cadena de texto EXACTA que aparezca en el documento (máximo 150 caracteres).
- No incluyas correcciones que no estén respaldadas por el texto del documento.
- Prioriza las inconsistencias de datos sobre los errores ortográficos menores.
- Si no hay nada que corregir, responde exactamente: []

Texto del documento:
{doc_text[:8000]}
"""
        client = get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Eres un revisor de documentos legales. Solo devuelves JSON válido."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=2000
        )

        content = response.choices[0].message.content.strip()
        # Limpiar markdown si el modelo lo incluye
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]

        corrections = json.loads(content.strip())
        if not isinstance(corrections, list):
            corrections = []

    except Exception as e:
        print(f"Error en review_document_with_ai (AI step): {e}")
        return {"corrections": 0, "details": [], "summary": "Error en revisión IA (API)."}

    if not corrections:
        return {"corrections": 0, "details": [], "summary": "Revisado por IA — sin correcciones necesarias."}

    # --- 3. Aplicar correcciones al .docx (find/replace preservando formato) ---
    applied = []
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(docx_path)

        def replace_in_paragraph(paragraph, original, corrected):
            """Reemplaza texto en un párrafo preservando el formato de runs."""
            full_text = paragraph.text
            if original not in full_text:
                return False
            # Reconstruir los runs para aplicar el reemplazo
            # Estrategia: concatenar texto de runs, detectar posición, redistribuir
            for run in paragraph.runs:
                if original in run.text:
                    run.text = run.text.replace(original, corrected, 1)
                    return True
            # Si el original cruza múltiples runs, combinar en el primero y limpiar los demás
            combined = ""
            run_texts = [r.text for r in paragraph.runs]
            combined = "".join(run_texts)
            if original in combined:
                new_combined = combined.replace(original, corrected, 1)
                if paragraph.runs:
                    paragraph.runs[0].text = new_combined
                    for r in paragraph.runs[1:]:
                        r.text = ""
                return True
            return False

        def replace_in_table(table, original, corrected):
            count = 0
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_in_paragraph(para, original, corrected):
                            count += 1
            return count

        for correction in corrections:
            original = correction.get("original", "")
            corrected = correction.get("corrected", "")
            if not original or not corrected or original == corrected:
                continue

            replaced = False
            # Párrafos del cuerpo principal
            for para in doc.paragraphs:
                if replace_in_paragraph(para, original, corrected):
                    replaced = True
                    break  # Aplicar solo la primera ocurrencia por par

            # Tablas del cuerpo principal
            if not replaced:
                for table in doc.tables:
                    if replace_in_table(table, original, corrected) > 0:
                        replaced = True
                        break

            if replaced:
                applied.append({"original": original, "corrected": corrected})

        if applied:
            doc.save(docx_path)

    except Exception as e:
        print(f"Error en review_document_with_ai (apply step): {e}")
        import traceback
        traceback.print_exc()
        # Si falla al aplicar, retornamos las correcciones sugeridas de todos modos
        n = len(corrections)
        return {"corrections": 0, "details": [], "summary": f"Revisión IA completada pero no se pudo aplicar al archivo ({str(e)})."}

    n = len(applied)
    if n == 0:
        summary = "Revisado por IA — sin correcciones aplicadas."
    elif n == 1:
        summary = "1 corrección de IA aplicada."
    else:
        summary = f"{n} correcciones de IA aplicadas."

    return {"corrections": n, "details": applied, "summary": summary}


def detect_variable_groups(variables: list) -> dict:
    """
    Auto-detecta grupos de variables numeradas.

    Patrones soportados:
      - {grupo}_{N}_{campo}  → ej: pago_1_monto, pago_2_monto
      - {grupo}_{N}          → ej: servicio_1, servicio_2

    Retorna:
      {
        "simple": ["cliente", "fecha"],
        "groups": [{"name": "pago", "label": "Pagos", "fields": ["monto", "concepto"]}]
      }
    """
    import re

    # Pattern: group_number_field  OR  group_number
    pattern_full = re.compile(r'^([a-z][a-z0-9]*(?:_[a-z][a-z0-9]*)*)_(\d+)_([a-z][a-z0-9_]*)$')
    pattern_simple = re.compile(r'^([a-z][a-z0-9]*(?:_[a-z][a-z0-9]*)*)_(\d+)$')

    groups_map = {}   # name -> set of fields
    group_order = []  # preserve first-seen order
    simple = []

    for var in variables:
        m = pattern_full.match(var)
        if m:
            group_name, _, field_name = m.group(1), m.group(2), m.group(3)
            if group_name not in groups_map:
                groups_map[group_name] = []
                group_order.append(group_name)
            if field_name not in groups_map[group_name]:
                groups_map[group_name].append(field_name)
            continue

        m2 = pattern_simple.match(var)
        if m2:
            group_name = m2.group(1)
            if group_name not in groups_map:
                groups_map[group_name] = []
                group_order.append(group_name)
            continue

        simple.append(var)

    def to_label(name: str) -> str:
        """Convierte 'pago' → 'Pagos', 'servicio' → 'Servicios'"""
        word = name.replace('_', ' ').strip().capitalize()
        # simple pluralización en español
        if word.endswith('o'):
            return word + 's'
        if word.endswith(('a', 'e')):
            return word + 's'
        return word + 'es'

    def field_label(name: str) -> str:
        return name.replace('_', ' ').capitalize()

    groups = [
        {
            "name": g,
            "label": to_label(g),
            "fields": [{"name": f, "label": field_label(f)} for f in groups_map[g]]
        }
        for g in group_order
    ]

    return {"simple": simple, "groups": groups}

def convert_to_pdf(docx_path: str) -> str:
    """
    Converts a .docx file to .pdf.
    - Windows/macOS: uses docx2pdf (requires MS Word).
    - Linux: uses libreoffice (headless).
    Returns the path to the generated PDF.
    """
    import sys
    import subprocess
    import platform
    import os
    
    pdf_path = docx_path.replace(".docx", ".pdf")
    
    try:
        if platform.system() == "Linux":
            # Linux: Use LibreOffice
            # libreoffice --headless --convert-to pdf <file> --outdir <dir>
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                docx_path,
                "--outdir",
                output_dir
            ]
            print(f"Converting to PDF using LibreOffice: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"LibreOffice failed: {result.stderr}")
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                
            # LibreOffice might change the filename case or something, but usually it keeps it.
            # It blindly replaces extension.
            
        else:
            # Windows/Mac: Use docx2pdf (requires MS Word installed)
            from docx2pdf import convert
            print(f"Converting {docx_path} to PDF using docx2pdf...")
            convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            raise Exception("PDF file was not created")
            
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return None


def extract_indexed_paragraphs(file_path: str) -> str:
    """Extrae párrafos numerados por índice para análisis de condicionales."""
    try:
        doc = Document(file_path)
        indexed = []
        for i, p in enumerate(doc.paragraphs):
            if len(p.text.strip()) > 5:
                indexed.append({"id": i, "text": p.text.strip()})
        return json.dumps(indexed, ensure_ascii=False)
    except Exception as e:
        print(f"Error extract_indexed_paragraphs: {e}")
        return "[]"


def generate_conditional_blocks_proposal(indexed_text_json: str, user_prompt: str = None) -> list:
    """Genera propuesta de bloques condicionales usando OpenAI."""
    if user_prompt and user_prompt.strip():
        extra_instructions = (
            f"\n\nINSTRUCCIONES DEL USUARIO (MÁXIMA PRIORIDAD — "
            f"guíate principalmente por esto para identificar los bloques):\n{user_prompt}"
        )
    else:
        extra_instructions = (
            "\n\nSin instrucciones específicas: identifica de forma general secciones, "
            "cláusulas, artículos, etapas, servicios, productos, condiciones o cualquier "
            "otro bloque lógico autónomo que un usuario podría querer activar o desactivar."
        )

    prompt = (
        "Eres un experto en estructuración y automatización de documentos de cualquier tipo "
        "(contratos, propuestas, manuales, pólizas, presupuestos, etc.).\n\n"
        "Recibirás los párrafos de un documento en formato JSON con \"id\" y \"text\". "
        "Tu tarea es identificar bloques de párrafos que representan secciones OPCIONALES: "
        "elementos que un usuario podría querer incluir o excluir al generar el documento final."
        f"{extra_instructions}\n\n"
        "REGLAS ESTRICTAS:\n"
        "1. 'variable_name': nombre corto en código, minúsculas, sin acentos, con guiones bajos. "
        "Ej: 'incluir_garantia', 'mostrar_anexo_tecnico'.\n"
        "2. 'label': pregunta clara y humana en español. Ej: '¿Desea incluir la garantía de calidad?'.\n"
        "3. 'start_index': el \"id\" exacto del PRIMER párrafo del bloque.\n"
        "4. 'end_index': el \"id\" exacto del ÚLTIMO párrafo del bloque (inclusive).\n"
        "5. Los bloques NO pueden solaparse.\n"
        "6. Agrupa siempre el título de la sección con su contenido en el mismo bloque.\n"
        "7. NO inventes bloques que no existen en el texto.\n\n"
        "Devuelve SOLO JSON válido con este esquema:\n"
        "{\"conditional_blocks\": [{\"variable_name\": \"incluir_seccion_x\", "
        "\"label\": \"¿Desea incluir la sección X?\", \"start_index\": 5, \"end_index\": 10}]}\n\n"
        f"Párrafos del documento (JSON indexado):\n{indexed_text_json[:15000]}"
    )
    try:
        client = get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Devuelve exclusivamente JSON válido según el esquema."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        content = response.choices[0].message.content.strip()
        data = json.loads(content)
        return data.get("conditional_blocks", [])
    except Exception as e:
        print(f"Error en generate_conditional_blocks_proposal: {e}")
        return []


def inject_conditions_to_docx(file_path: str, blocks: list) -> str:
    """Modifica el documento inyectando {%p if %} y {%p endif %} por índice de párrafo."""
    try:
        doc = Document(file_path)
        # Procesar de atrás hacia adelante para no alterar los índices al insertar
        sorted_blocks = sorted(blocks, key=lambda x: x['start_index'], reverse=True)

        for block in sorted_blocks:
            var_name = block.get('variable_name')
            start_id = block.get('start_index')
            end_id = block.get('end_index')

            if start_id < 0 or end_id >= len(doc.paragraphs):
                continue

            if end_id + 1 < len(doc.paragraphs):
                doc.paragraphs[end_id + 1].insert_paragraph_before('{%p endif %}')
            else:
                doc.add_paragraph('{%p endif %}')

            doc.paragraphs[start_id].insert_paragraph_before(
                f'{{%p if {var_name} %}}'
            )

        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"Error en inject_conditions_to_docx: {e}")
        return file_path
