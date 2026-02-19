import requests
import sys
import os

def test_generate_document():
    # 1. Login to get token
    login_url = "http://127.0.0.1:8000/api/v1/login/access-token"
    login_payload = {
        "username": "admin@example.com",
        "password": "admin"
    }
    
    try:
        print("Logging in...")
        login_response = requests.post(login_url, data=login_payload)
        if login_response.status_code != 200:
            print(f"Login failed: {login_response.text}")
            return
        
        token = login_response.json().get("access_token")
        headers = {"Authorization": f"Bearer {token}"}
        print("Login successful.")

        # 2. Upload a template (if not exists, or just ensure we have one)
        # For this test, we assume a template exists or we create a dummy one if we had a create_template function accessible here easily.
        # But we can list templates first.
        templates_url = "http://127.0.0.1:8000/api/v1/templates/"
        templates_res = requests.get(templates_url, headers=headers)
        templates = templates_res.json()
        
        template_id = None
        if templates:
            template_id = templates[0]["id"]
            print(f"Using existing template ID: {template_id}")
        else:
            print("No templates found. Creating and uploading a dummy template...")
            # Create a simple .docx file
            from docx import Document
            doc = Document()
            doc.add_paragraph("Hello {{nombre_cliente}}")
            doc.add_paragraph("Address: {{direccion}}")
            doc.add_paragraph("Amount: {{monto}}")
            dummy_filename = "dummy_template.docx"
            doc.save(dummy_filename)
            
            # Upload template
            # Note: The endpoint expects 'file' and 'title' as form data
            upload_url = "http://127.0.0.1:8000/api/v1/templates/"
            
            # Prepare the file and data
            with open(dummy_filename, 'rb') as f:
                files = {'file': (dummy_filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                data = {'title': 'Dummy Template'}
                
                upload_res = requests.post(upload_url, headers=headers, files=files, data=data)
            
            if upload_res.status_code == 200:
                template_data = upload_res.json()
                template_id = template_data["id"]
                print(f"Template uploaded successfully. ID: {template_id}")
                # Cleanup
                import time
                time.sleep(1) # wait for file handle release maybe?
                # os.remove(dummy_filename) # keep for inspection if needed
            else:
                 print(f"Failed to upload template: {upload_res.status_code} - {upload_res.text}")
                 return

        # 3. Generate Document
        generate_url = "http://127.0.0.1:8000/api/v1/documents/"
        payload = {
            "title": "Test Document Generation",
            "template_id": template_id,
            "variables": {
                "nombre_cliente": "Juan Perez",
                "direccion": "Av. Siempre Viva 123",
                "monto": "1000.00"
            }
        }
        
        print(f"Generating document with payload: {payload}")
        gen_response = requests.post(generate_url, json=payload, headers=headers)
        
        if gen_response.status_code == 200:
            doc_data = gen_response.json()
            print("Document generated successfully!")
            print(f"ID: {doc_data.get('id')}")
            print(f"Path: {doc_data.get('generated_file_path')}")
        else:
             print(f"Generation failed: {gen_response.status_code} - {gen_response.text}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    test_generate_document()
