import os
import requests
import docx

BASE = 'http://localhost:8001/api/v1'

def main():
    # 1. Login
    r = requests.post(f'{BASE}/login/access-token', data={'username':'contacto@simplaw.co','password':'Simplaw2026!'})
    token = r.json().get('access_token')
    headers = {'Authorization': f'Bearer {token}'}

    # 2. Create template with dynamic list variables
    doc = docx.Document()
    doc.add_heading('Contrato de Servicios - Prueba Dinámica', 0)
    doc.add_paragraph('Cliente: {{cliente}}')
    doc.add_paragraph('Primer servicio específico: {{servicio_1_nombre}} (Precio: {{servicio_1_precio}})')
    doc.add_paragraph('Segundo servicio específico: {{servicio_2_nombre}} (Precio: {{servicio_2_precio}})')
    doc.add_paragraph('Tercer servicio específico: {{servicio_3_nombre}} (Precio: {{servicio_3_precio}})')
    doc.add_heading('Lista completa de servicios:', level=2)
    doc.add_paragraph('{{lista_servicio}}')
    doc.add_paragraph('Total a pagar: {{precio_total_servicio}}')
    
    template_path = 'dynamic_test_template.docx'
    doc.save(template_path)

    # 3. Upload Template
    with open(template_path, 'rb') as f:
        files = {'file': ('dynamic_test_template.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'title': 'Dynamic Test Template'}
        r_upload = requests.post(f'{BASE}/templates/', files=files, data=data, headers=headers)
    
    if r_upload.status_code != 200:
        print("Upload failed:", r_upload.text)
        return
    template_id = r_upload.json()['id']
    print(f"Uploaded template ID: {template_id}")

    # 4. Update Schema for template (to simulate dynamic list configuration)
    schema = [
        {
            "type": "numbered_elements",
            "name": "servicio",
            "label": "Servicios Contratados",
            "fields": [
                {"name": "nombre", "label": "Nombre del Servicio", "type": "text"},
                {"name": "precio", "label": "Precio", "type": "number"}
            ]
        }
    ]
    r_schema = requests.put(f'{BASE}/templates/{template_id}', json={'variables_schema': schema}, headers=headers)
    if r_schema.status_code != 200:
        print("Schema update failed:", r_schema.text)
        return

    # 5. Generate Document
    payload = {
        'title': 'Test Documento Dinámico',
        'template_id': template_id,
        'variables': {
            'cliente': 'Empresa Ficticia SA',
            'servicio_items': [
                {'nombre': 'Desarrollo Web', 'precio': '3000'},
                {'nombre': 'Hosting Anual', 'precio': '150'},
                {'nombre': 'Mantenimiento', 'precio': '500'}
            ]
        }
    }
    r_gen = requests.post(f'{BASE}/documents/', json=payload, headers=headers)
    if r_gen.status_code != 200:
        print("Generation failed:", r_gen.json())
        return
    
    doc_data = r_gen.json()
    generated_path = doc_data['generated_file_path']
    print(f"Generated document path: {generated_path}")

    # 6. Verify Content
    gen_doc = docx.Document(generated_path)
    content = '[PARAGRAPHS]\n' + '\n'.join([p.text for p in gen_doc.paragraphs if p.text.strip()])
    print(content)
    
    # Assertions
    assert "Empresa Ficticia SA" in content
    assert "Desarrollo Web" in content
    assert "3000" in content
    assert "Hosting Anual" in content
    assert "Mantenimiento" in content
    # Total sum 3000+150+500 = 3650
    assert "3,650.00" in content

    print("\nSUCCESS: E2E dynamic list generation works flawlessly.")

if __name__ == '__main__':
    main()
