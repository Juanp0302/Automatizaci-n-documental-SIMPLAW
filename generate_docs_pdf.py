
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import time

def parse_markdown_to_docx(md_path, output_docx):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line = line.rstrip()
        
        # Code Lock
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[2:], level=3)
        
        # List Items (Bullets)
        elif line.strip().startswith('* ') or line.strip().startswith('- '):
            indent = len(line) - len(line.lstrip())
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 4 + 1))
        
        # Numbered List
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # Horizontal Rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50) # simple separator

        # Simple Paragraph
        elif line.strip():
            # Basic Bold parsing (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            # Empty line
            pass

    doc.save(output_docx)
    print(f"Created DOCX: {output_docx}")

def main():
    base_dir = r'c:\Users\Usuario\.gemini\antigravity\brain\0c128f6b-ffa4-4167-a0c4-212e1935eeb2'
    docs = [
        ('USER_MANUAL.md', 'USER_MANUAL.docx', 'USER_MANUAL.pdf'),
        ('TECHNICAL_DOCS.md', 'TECHNICAL_DOCS.docx', 'TECHNICAL_DOCS.pdf')
    ]

    for md, docx, pdf in docs:
        md_path = os.path.join(base_dir, md)
        docx_path = os.path.join(base_dir, docx)
        pdf_path = os.path.join(base_dir, pdf)

        if os.path.exists(md_path):
            print(f"Processing {md}...")
            # 1. Convert MD to DOCX
            try:
                parse_markdown_to_docx(md_path, docx_path)
            except Exception as e:
                print(f"Error creating DOCX for {md}: {e}")
                continue

            # 2. Convert DOCX to PDF
            print(f"Converting to PDF: {pdf}...")
            try:
                convert(docx_path, pdf_path)
                print(f"Success: {pdf_path}")
            except Exception as e:
                print(f"Error converting to PDF: {e}")
        else:
            print(f"File not found: {md_path}")

if __name__ == '__main__':
    main()
