
import os
import sys
from docx import Document

# Create a dummy docx
docx_path = os.path.abspath("test_doc.docx")
pdf_path = os.path.abspath("test_doc.pdf")

doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test document for PDF conversion.')
doc.save(docx_path)

print(f"Created {docx_path}")

try:
    import pythoncom
    pythoncom.CoInitialize()
    from docx2pdf import convert
    print("Converting to PDF...")
    convert(docx_path, pdf_path)
    print(f"Conversion successful: {os.path.exists(pdf_path)}")
except Exception as e:
    print(f"Conversion failed: {e}")
    import traceback
    traceback.print_exc()
finally:
    try:
        pythoncom.CoUninitialize()
    except:
        pass
