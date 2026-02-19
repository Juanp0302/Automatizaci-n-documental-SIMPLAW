
from docx import Document
from app.utils import extract_variables_from_docx
import os

def create_test_docx():
    doc = Document()
    doc.add_paragraph("First paragraph with {{ var1 }}")
    
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Table cell with {{ var2 }}"
    
    doc.add_paragraph("Second paragraph with {{ var3 }}")
    
    file_path = "test_order.docx"
    doc.save(file_path)
    return file_path

def test_extraction():
    file_path = create_test_docx()
    print(f"Created {file_path}")
    
    try:
        variables = extract_variables_from_docx(file_path)
        print(f"Extracted variables: {variables}")
        
        expected = ['var1', 'var2', 'var3']
        if variables == expected:
            print("SUCCESS: Variables are in correct order.")
        else:
            print(f"FAILURE: Expected {expected}, got {variables}")
            print("This confirms that paragraphs and tables are processed separately, losing their relative order.")
            
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

if __name__ == "__main__":
    test_extraction()
