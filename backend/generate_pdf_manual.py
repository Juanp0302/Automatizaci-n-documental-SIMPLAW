from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

try:
    from docx2pdf import convert
except ImportError:
    print("docx2pdf is not installed. PDF conversion will be skipped.")
    convert = None

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('Manual de Usuario: Creación de Plantillas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Este manual guía el proceso de creación, configuración y uso de plantillas en el sistema de Automatización Documental.')

    # Section 1: Preparar el Documento
    doc.add_heading('1. Preparar el Documento Word (.docx)', level=1)
    p = doc.add_paragraph('El sistema utiliza ')
    run = p.add_run('variables')
    run.bold = True
    p.add_run(' para identificar las partes dinámicas del documento.')
    
    doc.add_paragraph('Sintaxis: {{nombre_variable}}', style='Quote')
    
    doc.add_paragraph('Ejemplo: "En la ciudad de {{ciudad}}, a los {{dia}} días..."')
    
    doc.add_paragraph('Reglas:', style='List Bullet')
    doc.add_paragraph('Usa guiones bajos en lugar de espacios ({{nombre_completo}}).', style='List Bullet')
    doc.add_paragraph('El sistema respeta el formato (negrita, tamaño) de Word.', style='List Bullet')

    # Section 2: Subir Plantilla
    doc.add_heading('2. Subir la Plantilla', level=1)
    doc.add_paragraph('1. Ve a la sección Plantillas.')
    doc.add_paragraph('2. Haz clic en "Subir Plantilla".')
    
    # Add Screenshot if exists
    if os.path.exists("manual_assets/03_templates.png"):
        doc.add_picture("manual_assets/03_templates.png", width=Inches(6))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figura 1: Listado de Plantillas', style='Caption')

    # Section 3: Configurar Campos
    doc.add_heading('3. Configurar Campos', level=1)
    
    # Concepto Clave Box
    p = doc.add_paragraph()
    runner = p.add_run('Concepto Clave: El Word define QUÉ datos se necesitan. Tú defines CÓMO se piden (Menús, Fechas, Lógica).')
    runner.bold = True
    runner.italic = True
    
    doc.add_paragraph('Una vez subida la plantilla, utiliza el botón de Configuración (⚙️) para definir el comportamiento de cada variable.')
    
    if os.path.exists("manual_assets/05_template_config.png"):
        doc.add_picture("manual_assets/05_template_config.png", width=Inches(6))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figura 2: Configuración de Campos', style='Caption')

    doc.add_heading('Tipos de Campos', level=2)
    doc.add_paragraph('Texto: Para nombres o datos cortos.', style='List Bullet')
    doc.add_paragraph('Área de Texto: Para párrafos largos.', style='List Bullet')
    doc.add_paragraph('Lista Desplegable: Para opciones predefinidas.', style='List Bullet')
    doc.add_paragraph('Fecha: Selector de calendario.', style='List Bullet')

    doc.add_heading('Lógica Condicional', level=2)
    doc.add_paragraph('Puedes ocultar preguntas a menos que se cumpla una condición en otro campo (ej: Solo preguntar "Nombre Cónyuge" si "Estado Civil" es "Casado").')

    # Section 4: Generar Documento
    doc.add_heading('4. Generar Documento', level=1)
    doc.add_paragraph('Ve a "Nuevo Documento", selecciona tu plantilla y llena el formulario.')

    if os.path.exists("manual_assets/04_new_document.png"):
        doc.add_picture("manual_assets/04_new_document.png", width=Inches(6))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figura 3: Formulario de Generación', style='Caption')

    # Save
    output_docx = "MANUAL_PLANTILLAS.docx"
    doc.save(output_docx)
    print(f"Word document saved to {output_docx}")

    # Convert to PDF
    if convert:
        output_pdf = "MANUAL_PLANTILLAS.pdf"
        try:
            print("Converting to PDF...")
            convert(output_docx, output_pdf)
            print(f"PDF saved to {output_pdf}")
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            print("Please ensure Microsoft Word is installed to use docx2pdf.")

if __name__ == "__main__":
    create_manual()
