import os
import sys
from dotenv import load_dotenv
from openai import OpenAI
import docx

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_text(docx_path):
    doc = docx.Document(docx_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
            
    for table in doc.tables:
        text.append("--- INICIO TABLA ---")
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            text.append(" | ".join(row_data))
        text.append("--- FIN TABLA ---")
        
    return "\n".join(text)

def generate_template_proposal(doc_text):
    prompt = f"""
Eres un asistente legal experto en automatización de documentos usando la librería Python 'docxtpl' (que usa sintaxis Jinja2).
Se te proporcionará el texto extraído de un documento o contrato legal que ya fue diligenciado con datos estáticos de un cliente real.

Tu objetivo principal es:
1. Identificar en el texto cuáles son los DATOS VARIABLES que deberían cambiar cada vez que se cree un nuevo contrato (ej. Nombres, Cédulas, Fechas, Montos, Ciudades, Direcciones, Conceptos).
2. Devolver una lista exhaustiva en formato JSON de las variables (simples y repetitivas/tabulares) que se usarán para generar la plantilla.
3. Las variables simples deben tener un nombre en minúsculas tipo snake_case (ej. `nombre_cliente`, `cedula_cliente`, `fecha_firma`).
4. Si detectas TABLAS o listas de cosas repetitivas (ej. una tabla de pagos con Cuota, Fecha, Monto, o una lista de Servicios), debes crear un objeto tipo lista dinámica y asignar campos.

El formato JSON estricto que debes retornar es:
```json
{{
  "simple": [
    {{"name": "nombre_cliente", "label": "Nombre del Cliente", "type": "text"}},
    {{"name": "fecha_firma", "label": "Fecha de Firma", "type": "date"}},
    {{"name": "monto_total", "label": "Monto Total", "type": "text"}}
  ],
  "groups": [
    {{
      "name": "pago",
      "label": "Lista de Pagos",
      "fields": [
        {{"name": "numero_cuota", "label": "Número de Cuota", "type": "text"}},
        {{"name": "fecha_pago", "label": "Fecha de Pago", "type": "date"}},
        {{"name": "monto_cuota", "label": "Monto de la Cuota", "type": "text"}}
      ]
    }}
  ]
}}
```

Si no hay grupos o tablas repetitivas, "groups" debe quedar vacío `[]`.
Solo devuelve JSON válido. No devuelvas ningún texto extra, explicaciones ni etiquetas markdown.

Texto original del contrato:
{doc_text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Eres un asistente de automatización JSON."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1
    )
    
    return response.choices[0].message.content

def create_sample_docx(path):
    doc = docx.Document()
    doc.add_paragraph("CONTRATO DE PRESTACIÓN DE SERVICIOS")
    doc.add_paragraph("Entre los suscritos a saber: de una parte ABC ABOGADOS S.A.S. y de otra "
                      "parte JUAN PÉREZ con C.C. 1.234.567.890 de Bogotá, hemos convenido "
                      "celebrar el presente contrato el día 15 de marzo de 2024.")
    doc.add_paragraph("El precio pactado es de $5.000.000, los cuales se pagarán así:")
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cuota'
    hdr_cells[1].text = 'Fecha'
    hdr_cells[2].text = 'Monto'
    
    row1 = table.add_row().cells
    row1[0].text = '1'
    row1[1].text = '15 de abril de 2024'
    row1[2].text = '$2.500.000'
    
    row2 = table.add_row().cells
    row2[0].text = '2'
    row2[1].text = '15 de mayo de 2024'
    row2[2].text = '$2.500.000'
    
    doc.save(path)


if __name__ == "__main__":
    sample_path = "sample_test_ai.docx"
    print("Creando Word de prueba...")
    create_sample_docx(sample_path)
    
    print("Extrayendo texto...")
    text = extract_text(sample_path)
    print("Texto extraído:")
    print(text)
    print("-" * 40)
    
    print("Llamando a OpenAI...")
    try:
        result = generate_template_proposal(text)
        print("Resultado JSON:")
        print(result)
    except Exception as e:
        print("Error con OpenAI:", e)

    # Cleanup
    if os.path.exists(sample_path):
        os.remove(sample_path)
